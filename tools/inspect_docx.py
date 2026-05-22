
from docx import Document
import sys

def inspect(path):
    doc = Document(path)
    print("--- PARAGRAPHS ---")
    for p in doc.paragraphs:
        if p.text.strip():
            print(f"P: {p.text.strip()}")
            
    print("\n--- TABLES ---")
    for i, table in enumerate(doc.tables):
        print(f"Table {i}:")
        for r_idx, row in enumerate(table.rows):
            texts = [c.text.strip() for c in row.cells]
            # Unique texts (merged cells repeat)
            unique_texts = []
            seen = set()
            for t in texts:
                if t not in seen:
                    unique_texts.append(t)
                    seen.add(t)
            print(f"  Row {r_idx}: {unique_texts}")

if __name__ == "__main__":
    inspect("илова.docx")
