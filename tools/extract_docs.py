"""
Extract text from .doc and .docx files, convert Cyrillic to Latin, 
and prepare documents for the RAG knowledge base.
"""
import os
import sys
import re

# ─── Cyrillic → Latin conversion table for Uzbek ───
CYRILLIC_TO_LATIN = {
    # Uppercase special
    'Ў': "O'", 'Қ': 'Q', 'Ғ': "G'", 'Ҳ': 'H',
    'Ш': 'Sh', 'Ч': 'Ch', 'Ё': 'Yo', 'Ю': 'Yu', 'Я': 'Ya',
    'Ж': 'J', 'Ц': 'Ts', 'Щ': 'Sh',
    # Uppercase normal
    'А': 'A', 'Б': 'B', 'В': 'V', 'Г': 'G', 'Д': 'D',
    'Е': 'E', 'З': 'Z', 'И': 'I', 'Й': 'Y', 'К': 'K',
    'Л': 'L', 'М': 'M', 'Н': 'N', 'О': 'O', 'П': 'P',
    'Р': 'R', 'С': 'S', 'Т': 'T', 'У': 'U', 'Ф': 'F',
    'Х': 'X', 'Э': 'E', 'Ъ': "'",
    # Lowercase special
    'ў': "o'", 'қ': 'q', 'ғ': "g'", 'ҳ': 'h',
    'ш': 'sh', 'ч': 'ch', 'ё': 'yo', 'ю': 'yu', 'я': 'ya',
    'ж': 'j', 'ц': 'ts', 'щ': 'sh',
    # Lowercase normal  
    'а': 'a', 'б': 'b', 'в': 'v', 'г': 'g', 'д': 'd',
    'е': 'e', 'з': 'z', 'и': 'i', 'й': 'y', 'к': 'k',
    'л': 'l', 'м': 'm', 'н': 'n', 'о': 'o', 'п': 'p',
    'р': 'r', 'с': 's', 'т': 't', 'у': 'u', 'ф': 'f',
    'х': 'x', 'э': 'e', 'ъ': "'",
    # Signs to drop
    'Ь': '', 'ь': '',
}

def cyrillic_to_latin(text: str) -> str:
    """Convert Uzbek Cyrillic text to Latin script."""
    result = []
    i = 0
    while i < len(text):
        char = text[i]
        if char in CYRILLIC_TO_LATIN:
            result.append(CYRILLIC_TO_LATIN[char])
        else:
            result.append(char)
        i += 1
    return ''.join(result)


def extract_docx(filepath: str) -> str:
    """Extract text from .docx file."""
    import docx
    doc = docx.Document(filepath)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append('\t'.join(cells))
    
    return '\n'.join(paragraphs)


def extract_doc(filepath: str) -> str:
    """Extract text from .doc file using win32com (Word COM automation)."""
    import win32com.client
    import pythoncom
    
    pythoncom.CoInitialize()
    
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    
    abs_path = os.path.abspath(filepath)
    
    try:
        doc = word.Documents.Open(abs_path)
        text = doc.Content.Text
        doc.Close(False)
        return text
    finally:
        word.Quit()
        pythoncom.CoUninitialize()


def main():
    project_root = r'C:\Project\eco_voice_web'
    new_data_dir = os.path.join(project_root, 'new data')
    backend_dir = os.path.join(project_root, 'backend')
    
    # Files to process
    files_to_process = [
        # (filepath, needs_cyrillic_to_latin_conversion, output_name)
        (os.path.join(new_data_dir, '1036.doc'), False, '1036.txt'),
        (os.path.join(new_data_dir, '14.doc'), False, '14.txt'),
        (os.path.join(new_data_dir, '1.docx'), False, '1.txt'),
        (os.path.join(backend_dir, 'илова.docx'), True, 'ilova.txt'),
    ]
    
    output_dir = backend_dir
    
    for filepath, convert_cyrillic, output_name in files_to_process:
        print(f"\n{'='*60}")
        print(f"Processing: {filepath}")
        print(f"  Convert Cyrillic: {convert_cyrillic}")
        print(f"  Output: {output_name}")
        
        if not os.path.exists(filepath):
            print(f"  ❌ File not found: {filepath}")
            continue
        
        try:
            # Extract text
            if filepath.endswith('.docx'):
                text = extract_docx(filepath)
            elif filepath.endswith('.doc'):
                text = extract_doc(filepath)
            else:
                print(f"  ❌ Unknown format: {filepath}")
                continue
            
            # Convert Cyrillic to Latin if needed
            if convert_cyrillic:
                print(f"  🔄 Converting Cyrillic → Latin...")
                text = cyrillic_to_latin(text)
            
            # Clean up text
            # Remove excessive blank lines
            text = re.sub(r'\n{3,}', '\n\n', text)
            
            # Save
            output_path = os.path.join(output_dir, output_name)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            
            print(f"  ✅ Saved: {output_path}")
            print(f"  📊 Length: {len(text)} chars, {len(text.splitlines())} lines")
            print(f"  📊 Preview: {text[:200]}...")
            
        except Exception as e:
            print(f"  ❌ Error: {e}")
            import traceback
            traceback.print_exc()


if __name__ == '__main__':
    main()
